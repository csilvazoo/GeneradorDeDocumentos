import time
import re
import os
from selenium import webdriver
from selenium.webdriver.edge.service import Service as EdgeService
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.edge.options import Options
from docx import Document
from bs4 import BeautifulSoup

# Configuración de Edge
driver_path = None  # Si tienes el path del driver, colócalo aquí
options = Options()
options.add_argument("--start-maximized")
# options.add_argument("--headless")  # Descomenta si no quieres ver el navegador

driver = webdriver.Edge(options=options)

try:
    # Ruta actual para guardar el DOCX
    base_path = os.path.dirname(os.path.abspath(__file__))
    docx_path = os.path.join(base_path, "requerimientos_v3.docx")
    plantilla_path = os.path.join(base_path, "plantillaAFU2505.docx")

    # URL fija de la intranet
    url = "http://reportes03/reports/report/IyD/Gestion/Funcionalidad"
    driver.get(url)

    WebDriverWait(driver, 30).until(
        EC.presence_of_element_located((By.TAG_NAME, "body"))
    )
    print("Página cargada correctamente.")
    time.sleep(3)

    # Cambiar al iframe del reporte
    iframe = driver.find_element(By.TAG_NAME, "iframe")
    driver.switch_to.frame(iframe)
    print("Cambiado al iframe del reporte.")

    # Buscar el campo de número de funcionalidad y asignar el valor
    numero_funcionalidad = "11194"
    input_funcionalidad = WebDriverWait(driver, 30).until(
        EC.visibility_of_element_located((By.ID, "ReportViewerControl_ctl04_ctl03_txtValue"))
    )
    driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", input_funcionalidad)
    driver.execute_script("arguments[0].focus();", input_funcionalidad)
    input_funcionalidad.click()
    input_funcionalidad.clear()
    for char in numero_funcionalidad:
        input_funcionalidad.send_keys(char)
        time.sleep(0.1)
    print(f"Número de funcionalidad {numero_funcionalidad} asignado.")
    time.sleep(0.5)

    # Click en el botón "Ver informe"
    ver_informe_btn = WebDriverWait(driver, 20).until(
        EC.element_to_be_clickable((By.ID, "ReportViewerControl_ctl04_ctl00"))
    )
    ver_informe_btn.click()
    print("Botón 'Ver informe' clickeado dentro del iframe.")
    time.sleep(15)

    # Extraer todos los divs de la columna "Numero" y los de "Req. Cliente"
    divs = driver.find_elements(By.XPATH, "//div[contains(@style, 'width:20.22mm;min-width: 20.22mm;') or contains(@style, 'width:23.24mm;min-width: 23.24mm;')]")
    requerimiento_nros = []
    requerimiento_links = []
    after_req_cliente = False
    for div in divs:
        text = div.text.strip()
        # Detectar el marcador de inicio de requerimientos
        if text == "Req. Cliente":
            after_req_cliente = True
            continue
        if after_req_cliente:
            # Si es un número de requerimiento (no incidente)
            if text.isdigit() and len(text) < 7:  # Asumimos que los incidentes tienen 7 o más dígitos
                nro = text
                href = f"http://reportes03/reports/report/IyD/Gestion/Requerimiento?TipoRequerimiento=1&Numero={nro}"
                requerimiento_nros.append(nro)
                requerimiento_links.append(href)
            # Si es un incidente, lo salteamos
            elif text.isdigit() and len(text) >= 7:
                continue
            # Si encontramos otro "Req. Cliente", reiniciamos
            elif text == "Req. Cliente":
                continue
    print(f"Requerimientos encontrados: {requerimiento_nros}")

    document = Document(plantilla_path)
    
    def insert_paragraph_after(paragraph, text, style=None):
        new_paragraph = paragraph._parent.add_paragraph(text)
        paragraph._element.addnext(new_paragraph._element)
        if style:
            try:
                new_paragraph.style = style
            except KeyError:
                pass  # Si el estilo no existe, no aplicar nada extra
        return new_paragraph

    def insert_paragraph_with_custom_style(paragraph, text, style_name=None):
        new_paragraph = None
        if style_name:
            try:
                new_paragraph = paragraph._parent.add_paragraph(text, style=style_name)
            except KeyError:
                new_paragraph = paragraph._parent.add_paragraph(text)
        else:
            new_paragraph = paragraph._parent.add_paragraph(text)
        paragraph._element.addnext(new_paragraph._element)
        return new_paragraph

    # Buscar el párrafo con el título "Requerimientos" y guardar referencia
    req_paragraph = None
    for p in document.paragraphs:
        if p.text.strip().lower() == "requerimientos":
            req_paragraph = p
            break
    if req_paragraph is None:
        req_paragraph = document.paragraphs[-1]

    last_paragraph = req_paragraph
    for idx, nro in enumerate(requerimiento_nros):
        req_url = requerimiento_links[idx]
        driver.execute_script(f"window.open('{req_url}', '_blank');")
        driver.switch_to.window(driver.window_handles[-1])
        print(f"Requerimiento {nro} abierto en nueva pestaña.")
        time.sleep(5)
        html_req = driver.page_source
        soup = BeautifulSoup(html_req, 'html.parser')
        try:
            iframe_req = driver.find_element(By.TAG_NAME, "iframe")
            driver.switch_to.frame(iframe_req)
            print("Cambiado al iframe del requerimiento.")
            time.sleep(2)
            html_req = driver.page_source
            soup = BeautifulSoup(html_req, 'html.parser')
        except Exception as e:
            print(f"No se pudo cambiar al iframe del requerimiento: {e}")
            soup = BeautifulSoup(html_req, 'html.parser')
        finally:
            driver.switch_to.default_content()
        campos = {
            'Fecha alta': '',
            'Título': '',
            'Necesidad': '',
            'Implementación sugerida': '',
            'Cliente': '',
            'Relevamientos asociados': '',
            'Documento número': '',
            'Link': '',
            'Interacciones relacionadas': ''
        }
        documento_num = ''
        documento_link = ''
        incidente_num = ''
        incidente_link = ''
        for row in soup.find_all('tr'):
            celdas = row.find_all(['td', 'th'])
            if len(celdas) >= 2:
                label = celdas[0].get_text(strip=True)
                valor = celdas[1].get_text(strip=True)
                if 'fecha alta' in label.lower():
                    campos['Fecha alta'] = valor
                elif 'nombre' in label.lower():
                    campos['Título'] = valor
                elif 'necesidad' in label.lower():
                    campos['Necesidad'] = valor
                elif 'implem. sugerida' in label.lower():
                    campos['Implementación sugerida'] = valor
                elif 'cliente' in label.lower():
                    campos['Cliente'] += valor
                elif 'requerido por' in label.lower():
                    campos['Relevamientos asociados'] = valor
                elif 'documento' == label.lower():
                    doc_num_match = re.search(r'(\d+)', valor)
                    if doc_num_match:
                        documento_num = doc_num_match.group(1)
                        campos['Documento número'] = documento_num
                    else:
                        campos['Documento número'] = ''
                    link_tag = celdas[1].find('a', href=True)
                    if link_tag:
                        href = link_tag['href']
                        js_match = re.search(r"window.open\('([^']+)'", href)
                        if js_match:
                            documento_link = js_match.group(1)
                        else:
                            documento_link = href
                elif 'observación' in label.lower():
                    campos['Interacciones relacionadas'] = valor
        incidente_tag = soup.find('a', href=True, text=re.compile(r'\d{6,}'))
        if incidente_tag:
            incidente_num = incidente_tag.get_text(strip=True)
            href = incidente_tag['href']
            js_match = re.search(r"window.open\('([^']+)'", href)
            if js_match:
                incidente_link = js_match.group(1)
            else:
                incidente_link = href
        if campos['Cliente']:
            match = re.match(r"(\d+)(.*)", campos['Cliente'])
            if match:
                nro_cliente = match.group(1).strip()
                desc = match.group(2).strip()
                if desc:
                    campos['Cliente'] = f"{nro_cliente} - {desc}"
                else:
                    campos['Cliente'] = nro_cliente
        if documento_num:
            campos['Relevamientos asociados'] = ''
            campos['Documento número'] = documento_num
        else:
            campos['Relevamientos asociados'] = 'no hay'
            campos['Documento número'] = 'no hay'
            documento_link = ''
        for k in ['Fecha alta', 'Título', 'Necesidad', 'Implementación sugerida', 'Cliente']:
            if not campos[k]:
                campos[k] = 'no hay'
        # Insertar con estilos de la plantilla
        p = insert_paragraph_after(p, f"Requerimiento nro.{nro}", style="toa heading")
        p = insert_paragraph_after(p, f"Fecha alta: {campos['Fecha alta']}", style="List Bullet")
        p = insert_paragraph_after(p, f"Título: {campos['Título']}", style="List Bullet")
        p = insert_paragraph_after(p, f"Necesidad: {campos['Necesidad']}", style="List Bullet")
        p = insert_paragraph_after(p, f"Implementación sugerida: {campos['Implementación sugerida']}", style="List Bullet")
        p = insert_paragraph_after(p, f"Cliente: {campos['Cliente']}", style="List Bullet")
        p = insert_paragraph_after(p, f"Relevamientos asociados: {campos['Relevamientos asociados']}", style="List Bullet")
        p = insert_paragraph_after(p, f"Documento número: {campos['Documento número']}", style="List Bullet 2")
        if documento_link and documento_num:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            p = insert_paragraph_after(p, f"Link: ", style="List Bullet 2")
            r_id = document.part.relate_to(documento_link, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)
            new_run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            # Add the Hyperlink character style
            rStyle = OxmlElement('w:rStyle')
            rStyle.set(qn('w:val'), 'Hyperlink')
            rPr.append(rStyle)
            color = OxmlElement('w:color')
            color.set(qn('w:val'), '#009430')
            rPr.append(color)
            u = OxmlElement('w:u')
            u.set(qn('w:val'), 'single')
            rPr.append(u)
            new_run.append(rPr)
            t = OxmlElement('w:t')
            t.text = documento_num
            new_run.append(t)
            hyperlink.append(new_run)
            p._p.append(hyperlink)
        else:
            p = insert_paragraph_after(p, f"Link: no hay", style="List Bullet 2")
        if incidente_num and incidente_link:
            p = insert_paragraph_after(p, f"Interacciones relacionadas: ", style="List Bullet")
            p = insert_paragraph_after(p, f"Incidente: ", style="List Bullet 2")
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            r_id = document.part.relate_to(incidente_link, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)
            new_run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            # Add the Hyperlink character style
            rStyle = OxmlElement('w:rStyle')
            rStyle.set(qn('w:val'), 'Hyperlink')
            rPr.append(rStyle)
            color = OxmlElement('w:color')
            color.set(qn('w:val'), '#009430')
            rPr.append(color)
            u = OxmlElement('w:u')
            u.set(qn('w:val'), 'single')
            rPr.append(u)
            new_run.append(rPr)
            t = OxmlElement('w:t')
            t.text = incidente_num
            new_run.append(t)
            hyperlink.append(new_run)
            p._p.append(hyperlink)
        else:
            p = insert_paragraph_after(p, f"Interacciones relacionadas: no hay", style="List Bullet")
        last_paragraph = p

    # Eliminar el último bloque de requerimiento si está vacío (de atrás hacia adelante)
    def is_empty_req_block(paragraphs, start_idx, labels):
        end_idx = start_idx + 1
        while end_idx < len(paragraphs):
            text = paragraphs[end_idx].text.strip().lower()
            if text.startswith('requerimiento nro.') or text.startswith('funcionalidad nro.'):
                break
            end_idx += 1
        for i in range(start_idx, end_idx):
            t = paragraphs[i].text.strip().lower()
            if t and not any(t.startswith(label) for label in labels):
                return False, end_idx
            if ':' in t:
                val = t.split(':',1)[1].strip()
                if val and val != 'no hay':
                    return False, end_idx
        return True, end_idx

    labels = [
        'requerimiento nro.',
        'fecha alta',
        'título',
        'necesidad',
        'implementación sugerida',
        'cliente',
        'relevamientos asociados',
        'documento número',
        'link',
        'interacciones relacionadas',
        'incidente'
    ]
    for i in range(len(document.paragraphs)-1, -1, -1):
        t = document.paragraphs[i].text.strip().lower()
        if t.startswith('requerimiento nro.'):
            empty, end_idx = is_empty_req_block(document.paragraphs, i, labels)
            if empty:
                for j in range(end_idx-1, i-1, -1):
                    document.paragraphs[j]._element.getparent().remove(document.paragraphs[j]._element)
                break

    # --- AGREGAR DATOS DE LA FUNCIONALIDAD ---
    # Cambiar al iframe del reporte antes de extraer los datos de la funcionalidad
    driver.switch_to.window(driver.window_handles[0])  # Volver a la pestaña de Funcionalidad
    driver.switch_to.default_content()
    iframe = driver.find_element(By.TAG_NAME, "iframe")
    driver.switch_to.frame(iframe)
    html_func = driver.page_source
    soup_func = BeautifulSoup(html_func, 'html.parser')
    campos_func = {
        'Funcionalidad nro.': numero_funcionalidad,
        'Nombre': '',
        'Descripción': '',
        'Producto': '',
        'Equipo': '',
        'Fecha alta': ''
    }
    # Lógica robusta para extraer los campos de funcionalidad
    for row_idx, row in enumerate(soup_func.find_all('tr')):
        celdas = row.find_all(['td', 'th'])
        for idx, celda in enumerate(celdas):
            texto = celda.get_text(strip=True).lower().replace(':','').replace('á','a').replace('é','e').replace('í','i').replace('ó','o').replace('ú','u')
            # Nombre y Descripción funcionan bien con la celda siguiente
            if texto == 'nombre' and not campos_func['Nombre'] and idx+1 < len(celdas):
                campos_func['Nombre'] = celdas[idx+1].get_text(strip=True)
            elif texto == 'descripcion' and not campos_func['Descripción'] and idx+1 < len(celdas):
                campos_func['Descripción'] = celdas[idx+1].get_text(strip=True)
            # Producto: buscar en la celda siguiente, y si está vacía, buscar en la siguiente fila, misma columna
            elif texto in ['producto', 'productos'] and not campos_func['Producto']:
                valor = ''
                # 1. Celda siguiente
                if idx+1 < len(celdas):
                    valor = celdas[idx+1].get_text(strip=True)
                    # Si la celda tiene una tabla, buscar el primer texto no vacío
                    if not valor:
                        inner_table = celdas[idx+1].find('table')
                        if inner_table:
                            valor = inner_table.get_text(strip=True)
                # 2. Si sigue vacío, buscar en la siguiente fila, misma columna
                if not valor:
                    filas = soup_func.find_all('tr')
                    if row_idx+1 < len(filas):
                        next_row = filas[row_idx+1]
                        next_celdas = next_row.find_all(['td', 'th'])
                        if idx < len(next_celdas):
                            valor = next_celdas[idx].get_text(strip=True)
                            if not valor:
                                inner_table = next_celdas[idx].find('table')
                                if inner_table:
                                    valor = inner_table.get_text(strip=True)
                campos_func['Producto'] = valor
            # Equipo: igual que antes
            elif texto == 'equipo' and not campos_func['Equipo'] and idx+1 < len(celdas):
                campos_func['Equipo'] = celdas[idx+1].get_text(strip=True)
            # Fecha alta: buscar en la celda siguiente, y si está vacía, buscar en la siguiente fila, misma columna
            elif texto == 'fecha alta' and not campos_func['Fecha alta']:
                valor = ''
                if idx+1 < len(celdas):
                    valor = celdas[idx+1].get_text(strip=True)
                    if not valor:
                        inner_table = celdas[idx+1].find('table')
                        if inner_table:
                            valor = inner_table.get_text(strip=True)
                if not valor:
                    filas = soup_func.find_all('tr')
                    if row_idx+1 < len(filas):
                        next_row = filas[row_idx+1]
                        next_celdas = next_row.find_all(['td', 'th'])
                        if idx < len(next_celdas):
                            valor = next_celdas[idx].get_text(strip=True)
                            if not valor:
                                inner_table = next_celdas[idx].find('table')
                                if inner_table:
                                    valor = inner_table.get_text(strip=True)
                campos_func['Fecha alta'] = valor
    # Fallback: si algún campo sigue vacío, dejar 'no hay'
    for k in campos_func:
        if not campos_func[k]:
            campos_func[k] = 'no hay'
    # Insertar los datos de la funcionalidad después del último requerimiento real y antes de 'Notas aclaratorias' o al final si no existe
    insert_idx = None
    for i, p in enumerate(document.paragraphs):
        if p.text.strip().lower().startswith('funcionalidad nro.'):
            insert_idx = i
            break
    if insert_idx is not None:
        end_idx = insert_idx + 1
        while end_idx < len(document.paragraphs):
            t = document.paragraphs[end_idx].text.strip().lower()
            if t.startswith('•') or any(x in t for x in ['nombre', 'descrip', 'producto', 'equipo', 'fecha alta']):
                end_idx += 1
            else:
                break
        for j in range(end_idx-1, insert_idx-1, -1):
            document.paragraphs[j]._element.getparent().remove(document.paragraphs[j]._element)
        p = document.paragraphs[insert_idx-1] if insert_idx > 0 else document.paragraphs[0]
    else:
        p = last_paragraph
    p = insert_paragraph_after(p, f"Funcionalidad nro.{campos_func['Funcionalidad nro.']}", style="Heading 2")
    p = insert_paragraph_after(p, f"Nombre: {campos_func['Nombre']}", style="List Bullet")
    p = insert_paragraph_after(p, f"Descripción: {campos_func['Descripción']}", style="List Bullet")
    p = insert_paragraph_after(p, f"Producto: {campos_func['Producto']}", style="List Bullet")
    p = insert_paragraph_after(p, f"Equipo: {campos_func['Equipo']}", style="List Bullet")
    p = insert_paragraph_after(p, f"Fecha alta: {campos_func['Fecha alta']}", style="List Bullet")
    last_paragraph = p

    document.save(docx_path)
    print(f"Documento Word guardado en {docx_path}")
    print("Proceso finalizado. Cierra el navegador.")
finally:
    driver.quit()
