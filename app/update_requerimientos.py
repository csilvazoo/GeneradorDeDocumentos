import os
import copy
from docx import Document
from app.docx_helpers import insert_paragraph_after, is_empty_req_block, add_hyperlink
from app.requerimientos import extraer_requerimiento
from app.funcionalidad import extraer_funcionalidad
from app.selenium_helpers import cambiar_iframe, abrir_pestania
from app.utils import log_to_queue, validar_numero_funcionalidad
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import time

def update_requerimientos(numero_funcionalidad, docx_path, driver, log_queue):
    log = lambda msg: log_to_queue(log_queue, msg)
    # Abrir el documento seleccionado por el usuario
    document = Document(docx_path)
    # Obtener los requerimientos ya presentes en el documento
    req_nros_en_doc = set()
    for p in document.paragraphs:
        t = p.text.strip().lower()
        if t.startswith('requerimiento nro.'):
            nro = t.replace('requerimiento nro.', '').strip()
            if nro.isdigit():
                req_nros_en_doc.add(nro)
    # Navegar y obtener los requerimientos actuales de la intranet
    url = "http://reportes03/reports/report/IyD/Gestion/Funcionalidad"
    driver.get(url)
    WebDriverWait(driver, 30).until(EC.presence_of_element_located((By.TAG_NAME, "body")))
    log("Página cargada correctamente.")
    time.sleep(3)
    cambiar_iframe(driver, By.TAG_NAME, "iframe")
    log("Cambiado al iframe del reporte.")
    input_funcionalidad = WebDriverWait(driver, 30).until(
        EC.visibility_of_element_located((By.ID, "ReportViewerControl_ctl04_ctl03_txtValue"))
    )
    driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", input_funcionalidad)
    driver.execute_script("arguments[0].focus();", input_funcionalidad)
    input_funcionalidad.click()
    input_funcionalidad.clear()
    for char in numero_funcionalidad:
        input_funcionalidad.send_keys(char)
        time.sleep(0.1)
    log(f"Número de funcionalidad {numero_funcionalidad} asignado.")
    time.sleep(0.5)
    ver_informe_btn = WebDriverWait(driver, 20).until(
        EC.element_to_be_clickable((By.ID, "ReportViewerControl_ctl04_ctl00"))
    )
    ver_informe_btn.click()
    log("Botón 'Ver informe' clickeado dentro del iframe.")
    time.sleep(15)
    divs = driver.find_elements(By.XPATH, "//div[contains(@style, 'width:20.22mm;min-width: 20.22mm;') or contains(@style, 'width:23.24mm;min-width: 23.24mm;')]")
    requerimiento_nros = []
    requerimiento_links = []
    after_req_cliente = False
    for div in divs:
        text = div.text.strip()
        if text == "Req. Cliente":
            after_req_cliente = True
            continue
        if after_req_cliente:
            if text.isdigit() and len(text) < 7:
                nro = text
                href = f"http://reportes03/reports/report/IyD/Gestion/Requerimiento?TipoRequerimiento=1&Numero={nro}"
                requerimiento_nros.append(nro)
                requerimiento_links.append(href)
            elif text.isdigit() and len(text) >= 7:
                continue
            elif text == "Req. Cliente":
                continue
    log(f"Requerimientos encontrados: {requerimiento_nros}")
    log(f"Requerimientos existentes en el documento:")
    for idx, nro in enumerate(sorted(req_nros_en_doc), 1):
        log(f"  {idx}. {nro}")
    nuevos_reqs = [nro for nro in requerimiento_nros if nro not in req_nros_en_doc]
    log(f"Requerimientos a agregar:")
    for idx, nro in enumerate(nuevos_reqs, 1):
        log(f"  {idx}. {nro}")
    if not nuevos_reqs:
        log("No hay nuevos requerimientos para agregar al documento.")
        log("Proceso finalizado.")
        return False
    # 1. Detectar el bloque de 'Funcionalidad nro.' y sus detalles
    func_start_idx = None
    func_end_idx = None
    
    # Buscar el párrafo que contiene 'Funcionalidad nro.'
    for i, p in enumerate(document.paragraphs):
        if p.text.strip().lower().startswith('funcionalidad nro.'):
            func_start_idx = i
            break
    
    if func_start_idx is not None:
        # Encontrar el final del bloque de funcionalidad
        func_end_idx = func_start_idx + 1
        while func_end_idx < len(document.paragraphs):
            t = document.paragraphs[func_end_idx].text.strip().lower()
            if t.startswith('requerimiento nro.'):
                break
            func_end_idx += 1
        
        # Obtener el párrafo de inicio y fin del bloque
        func_start_element = document.paragraphs[func_start_idx]._element
        if func_end_idx < len(document.paragraphs):
            func_end_element = document.paragraphs[func_end_idx]._element
        else:
            func_end_element = None
        
        # Extraer TODOS los elementos entre el inicio y fin (incluyendo tablas)
        func_block_elements = []
        current_element = func_start_element
        
        while current_element is not None:
            func_block_elements.append(current_element)
            if current_element == func_end_element:
                break
            current_element = current_element.getnext()
            if current_element == func_end_element:
                break
        
        # Eliminar los elementos del bloque de funcionalidad del documento
        for element in func_block_elements:
            element.getparent().remove(element)
        
        # Encontrar dónde insertar los nuevos requerimientos
        insert_after_paragraph = document.paragraphs[func_start_idx-1] if func_start_idx > 0 else None
    else:
        func_block_elements = []
        insert_after_paragraph = document.paragraphs[-1]
    # 2. Insertar los nuevos requerimientos
    last_p = insert_after_paragraph
    for nro in nuevos_reqs:
        idx = requerimiento_nros.index(nro)
        req_url = requerimiento_links[idx]
        abrir_pestania(driver, req_url, log, nro_req=nro)
        campos, documento_num, documento_link, incidente_num, incidente_link = extraer_requerimiento(driver, req_url, log)
        p = insert_paragraph_after(last_p, f"Requerimiento nro.{nro}", style="toa heading")
        p = insert_paragraph_after(p, f"Fecha alta: {campos['Fecha alta']}", style="List Bullet")
        p = insert_paragraph_after(p, f"Título: {campos['Título']}", style="List Bullet")
        p = insert_paragraph_after(p, f"Necesidad: {campos['Necesidad']}", style="List Bullet")
        p = insert_paragraph_after(p, f"Implementación sugerida: {campos['Implementación sugerida']}", style="List Bullet")
        p = insert_paragraph_after(p, f"Cliente: {campos['Cliente']}", style="List Bullet")
        p = insert_paragraph_after(p, f"Relevamientos asociados: {campos['Relevamientos asociados']}", style="List Bullet")
        p = insert_paragraph_after(p, f"Documento número: {campos['Documento número']}", style="List Bullet 2")
        if documento_link and documento_num:
            p = insert_paragraph_after(p, f"Link: ", style="List Bullet 2")
            add_hyperlink(p, documento_link, documento_num, document)
        else:
            p = insert_paragraph_after(p, f"Link: no hay", style="List Bullet 2")
        if incidente_num and incidente_link:
            p = insert_paragraph_after(p, f"Interacciones relacionadas: ", style="List Bullet")
            p = insert_paragraph_after(p, f"Incidente: ", style="List Bullet 2")
            add_hyperlink(p, incidente_link, incidente_num, document)
        else:
            p = insert_paragraph_after(p, f"Interacciones relacionadas: no hay", style="List Bullet")
        last_p = p
    # 3. Volver a insertar el bloque de funcionalidad después del último requerimiento
    if func_block_elements:
        # Insertar los elementos XML preservados en orden (mantiene párrafos, tablas, imágenes, comentarios, etc.)
        current_insert_point = last_p._element
        
        for element in func_block_elements:
            # Hacer una copia profunda del elemento XML para preservar todo el contenido
            cloned_element = copy.deepcopy(element)
            # Insertar después del punto de inserción actual
            current_insert_point.addnext(cloned_element)
            # Actualizar el punto de inserción para el siguiente elemento
            current_insert_point = cloned_element
            
            # Si el elemento clonado es un párrafo, actualizar last_p
            if cloned_element.tag.endswith('}p'):  # Es un párrafo
                for p in document.paragraphs:
                    if p._element == cloned_element:
                        last_p = p
                        break
    document.save(docx_path)
    log(f"Documento actualizado con nuevos requerimientos: {nuevos_reqs}")
    log("Proceso finalizado.")
    return True