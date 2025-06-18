# Lógica de negocio para la generación y actualización de documentos de funcionalidad

# Aquí se moverán las funciones como run_script y helpers.

import sys
import os
import queue
import re
import time
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.edge.options import Options
from selenium.webdriver.edge.service import Service as EdgeService
from docx import Document
from bs4 import BeautifulSoup
from app.docx_helpers import insert_paragraph_after, is_empty_req_block, add_hyperlink
from app.requerimientos import extraer_requerimiento
from app.funcionalidad import extraer_funcionalidad
from app.selenium_helpers import cambiar_iframe, abrir_pestania
from app.utils import log_to_queue, validar_numero_funcionalidad

def run_script(numero_funcionalidad, ver_explorador, log_queue, docx_path):
    log = lambda msg: log_to_queue(log_queue, msg)
    if not validar_numero_funcionalidad(numero_funcionalidad):
        log("Número de funcionalidad inválido. Debe ser numérico.")
        return
    options = Options()
    options.add_argument("--start-maximized")
    if not ver_explorador:
        options.add_argument("--headless")
    if getattr(sys, 'frozen', False):
        driver_path = os.path.join(sys._MEIPASS, "msedgedriver.exe")
    else:
        driver_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "msedgedriver.exe")
    service = EdgeService(driver_path)
    driver = webdriver.Edge(service=service, options=options)
    try:
        base_path = os.path.dirname(os.path.abspath(__file__))
        plantilla_path = os.path.join(base_path, "..", "resources", "templates", "plantillaAFU2505.docx")
        plantilla_path = os.path.abspath(plantilla_path)
        url = "http://reportes03/reports/report/IyD/Gestion/Funcionalidad"
        driver.get(url)
        WebDriverWait(driver, 30).until(
            EC.presence_of_element_located((By.TAG_NAME, "body"))
        )
        log("Página cargada correctamente.")
        time.sleep(3)
        # Uso de helper para cambiar al iframe del reporte
        cambiar_iframe(driver, By.TAG_NAME, "iframe")
        log("Cambiado al iframe del reporte.")
        input_funcionalidad = WebDriverWait(driver, 30).until(
            EC.visibility_of_element_located((By.ID, "ReportViewerControl_ctl04_ctl03_txtValue"))
        )
        driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", input_funcionalidad)
        driver.execute_script("arguments[0].focus();", input_funcionalidad)
        input_funcionalidad.click()
        input_funcionalidad.clear()
        for char in numero_funcionalidad:
            input_funcionalidad.send_keys(char)
            time.sleep(0.1)
        log(f"Número de funcionalidad {numero_funcionalidad} asignado.")
        time.sleep(0.5)
        ver_informe_btn = WebDriverWait(driver, 20).until(
            EC.element_to_be_clickable((By.ID, "ReportViewerControl_ctl04_ctl00"))
        )
        ver_informe_btn.click()
        log("Botón 'Ver informe' clickeado dentro del iframe.")
        time.sleep(15)
        divs = driver.find_elements(By.XPATH, "//div[contains(@style, 'width:20.22mm;min-width: 20.22mm;') or contains(@style, 'width:23.24mm;min-width: 23.24mm;')]")
        requerimiento_nros = []
        requerimiento_links = []
        after_req_cliente = False
        for div in divs:
            text = div.text.strip()
            if text == "Req. Cliente":
                after_req_cliente = True
                continue
            if after_req_cliente:
                if text.isdigit() and len(text) < 7:
                    nro = text
                    href = f"http://reportes03/reports/report/IyD/Gestion/Requerimiento?TipoRequerimiento=1&Numero={nro}"
                    requerimiento_nros.append(nro)
                    requerimiento_links.append(href)
                elif text.isdigit() and len(text) >= 7:
                    continue
                elif text == "Req. Cliente":
                    continue
        log(f"Requerimientos encontrados: {requerimiento_nros}")
        if not requerimiento_nros:
            log("No existe funcionalidad o no tiene requerimientos asociados.")
            try:
                import tkinter
                from tkinter import messagebox
                root = tkinter.Tk()
                root.withdraw()
                messagebox.showerror("Error", "No existe funcionalidad o no tiene requerimientos asociados.")
                root.destroy()
            except Exception:
                pass
            log("Proceso finalizado.")
            return
        document = Document(plantilla_path)
        req_paragraph = None
        for p in document.paragraphs:
            if p.text.strip().lower() == "requerimientos":
                req_paragraph = p
                break
        if req_paragraph is None:
            req_paragraph = document.paragraphs[-1]
        last_paragraph = req_paragraph
        for idx, nro in enumerate(requerimiento_nros):
            req_url = requerimiento_links[idx]
            abrir_pestania(driver, req_url, log, nro_req=nro)
            campos, documento_num, documento_link, incidente_num, incidente_link = extraer_requerimiento(driver, req_url, log)
            p = insert_paragraph_after(last_paragraph, f"Requerimiento nro.{nro}", style="toa heading")
            p = insert_paragraph_after(p, f"Fecha alta: {campos['Fecha alta']}", style="List Bullet")
            p = insert_paragraph_after(p, f"Título: {campos['Título']}", style="List Bullet")
            p = insert_paragraph_after(p, f"Necesidad: {campos['Necesidad']}", style="List Bullet")
            p = insert_paragraph_after(p, f"Implementación sugerida: {campos['Implementación sugerida']}", style="List Bullet")
            p = insert_paragraph_after(p, f"Cliente: {campos['Cliente']}", style="List Bullet")
            p = insert_paragraph_after(p, f"Relevamientos asociados: {campos['Relevamientos asociados']}", style="List Bullet")
            p = insert_paragraph_after(p, f"Documento número: {campos['Documento número']}", style="List Bullet 2")
            if documento_link and documento_num:
                p = insert_paragraph_after(p, f"Link: ", style="List Bullet 2")
                add_hyperlink(p, documento_link, documento_num, document)
            else:
                p = insert_paragraph_after(p, f"Link: no hay", style="List Bullet 2")
            if incidente_num and incidente_link:
                p = insert_paragraph_after(p, f"Interacciones relacionadas: ", style="List Bullet")
                p = insert_paragraph_after(p, f"Incidente: ", style="List Bullet 2")
                add_hyperlink(p, incidente_link, incidente_num, document)
            else:
                p = insert_paragraph_after(p, f"Interacciones relacionadas: no hay", style="List Bullet")
            last_paragraph = p
        labels = [
            'requerimiento nro.',
            'fecha alta',
            'título',
            'necesidad',
            'implementación sugerida',
            'cliente',
            'relevamientos asociados',
            'documento número',
            'link',
            'interacciones relacionadas',
            'incidente'
        ]
        for i in range(len(document.paragraphs)-1, -1, -1):
            t = document.paragraphs[i].text.strip().lower()
            if t.startswith('requerimiento nro.'):
                empty, end_idx = is_empty_req_block(document.paragraphs, i, labels)
                if empty:
                    for j in range(end_idx-1, i-1, -1):
                        document.paragraphs[j]._element.getparent().remove(document.paragraphs[j]._element)
                    break
        # Cambiar a la ventana principal y al iframe usando helper
        driver.switch_to.window(driver.window_handles[0])
        driver.switch_to.default_content()
        cambiar_iframe(driver, By.TAG_NAME, "iframe")
        log("No hay más requerimientos asociados.")
        time.sleep(2)
        log("Cambiado al iframe de la funcionalidad.")
        # --- EXTRACCIÓN DE DATOS DE FUNCIONALIDAD MODULARIZADA ---
        html_func = driver.page_source
        campos_func = extraer_funcionalidad(html_func, numero_funcionalidad)
        insert_idx = None
        for i, p in enumerate(document.paragraphs):
            if p.text.strip().lower().startswith('funcionalidad nro.'):
                insert_idx = i
                break
        if insert_idx is not None:
            end_idx = insert_idx + 1
            while end_idx < len(document.paragraphs):
                t = document.paragraphs[end_idx].text.strip().lower()
                if t.startswith('•') or any(x in t for x in ['nombre', 'descrip', 'producto', 'equipo', 'fecha alta']):
                    end_idx += 1
                else:
                    break
            for j in range(end_idx-1, insert_idx-1, -1):
                document.paragraphs[j]._element.getparent().remove(document.paragraphs[j]._element)
            p = document.paragraphs[insert_idx-1] if insert_idx > 0 else document.paragraphs[0]
        else:
            p = last_paragraph
        p = insert_paragraph_after(p, f"Funcionalidad nro.{campos_func['Funcionalidad nro.']}", style="Heading 2")
        p = insert_paragraph_after(p, f"Nombre: {campos_func['Nombre']}", style="List Bullet")
        p = insert_paragraph_after(p, f"Descripción: {campos_func['Descripción']}", style="List Bullet")
        p = insert_paragraph_after(p, f"Producto: {campos_func['Producto']}", style="List Bullet")
        p = insert_paragraph_after(p, f"Equipo: {campos_func['Equipo']}", style="List Bullet")
        p = insert_paragraph_after(p, f"Fecha alta: {campos_func['Fecha alta']}", style="List Bullet")
        last_paragraph = p
        document.save(docx_path)
        log(f"Documento Word guardado en {docx_path}")
        log("Proceso finalizado.")
    except Exception as e:
        log(f"ERROR: {e}")
        log("Proceso finalizado.")
    finally:
        driver.quit()
