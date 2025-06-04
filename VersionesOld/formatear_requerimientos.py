from docx import Document
import re

# Abrir el docx generado previamente
doc = Document(r"c:\Users\csilva\Desktop\requerimientos_funcionalidad.docx")

# Crear nuevo docx para el formato solicitado
nuevo_doc = Document()

# Expresiones regulares para extraer campos
re_num = re.compile(r"(Req\.? Cliente|Requerimiento nro\.?)(\s*\d+)", re.IGNORECASE)
re_fecha = re.compile(r"(\d{1,2}/\d{1,2}/\d{4})")

req_actual = {}
reqs = []

for para in doc.paragraphs:
    text = para.text.strip()
    if not text:
        if req_actual:
            reqs.append(req_actual)
            req_actual = {}
        continue
    m_num = re_num.search(text)
    if m_num:
        if req_actual:
            reqs.append(req_actual)
            req_actual = {}
        req_actual['nro'] = m_num.group(2)
        continue
    m_fecha = re_fecha.search(text)
    if m_fecha:
        req_actual['fecha'] = m_fecha.group(1)
        continue
    if text.lower().startswith('título:'):
        req_actual['titulo'] = text.split(':',1)[1].strip()
        continue
    if text.lower().startswith('necesidad:'):
        req_actual['necesidad'] = text.split(':',1)[1].strip()
        continue
    if text.lower().startswith('implementación sugerida:'):
        req_actual['implementacion'] = text.split(':',1)[1].strip()
        continue
    if text.lower().startswith('cliente:'):
        req_actual['cliente'] = text.split(':',1)[1].strip()
        continue
    # Otros campos pueden agregarse aquí

if req_actual:
    reqs.append(req_actual)

# Escribir en el nuevo docx con el formato solicitado
for req in reqs:
    if 'nro' in req:
        nuevo_doc.add_paragraph(f"Requerimiento nro.{req['nro']}")
    if 'fecha' in req:
        nuevo_doc.add_paragraph(f"•\tFecha alta: {req['fecha']}")
    if 'titulo' in req:
        nuevo_doc.add_paragraph(f"•\tTítulo: {req['titulo']}")
    if 'necesidad' in req:
        nuevo_doc.add_paragraph(f"•\tNecesidad: {req['necesidad']}")
    if 'implementacion' in req:
        nuevo_doc.add_paragraph(f"•\tImplementación sugerida: {req['implementacion']}")
    if 'cliente' in req:
        nuevo_doc.add_paragraph(f"•\tCliente: {req['cliente']}")
    # Campos fijos para el formato solicitado
    nuevo_doc.add_paragraph(f"•\tRelevamientos asociados:")
    nuevo_doc.add_paragraph(f"•\tDocumento número:")
    nuevo_doc.add_paragraph(f"•\tLink:")
    nuevo_doc.add_paragraph(f"•\tInteracciones relacionadas: no hay")
    nuevo_doc.add_paragraph("")

nuevo_doc.save(r"c:\Users\csilva\Desktop\requerimientos_formato.docx")
print("Documento Word generado: requerimientos_formato.docx")
