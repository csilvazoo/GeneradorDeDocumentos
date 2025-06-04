import time
import re
import os
from selenium import webdriver
from selenium.webdriver.edge.service import Service as EdgeService
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.edge.options import Options
from docx import Document
from bs4 import BeautifulSoup

# Configuración de Edge
driver_path = None  # Si tienes el path del driver, colócalo aquí
options = Options()
options.add_argument("--start-maximized")
# options.add_argument("--headless")  # Descomenta si no quieres ver el navegador

driver = webdriver.Edge(options=options)

try:
    # Ruta actual para guardar el HTML y el DOCX
    base_path = os.path.dirname(os.path.abspath(__file__))
    html_path = os.path.join(base_path, "reporte_funcionalidad.html")
    docx_path = os.path.join(base_path, "requerimientos.docx")

    # URL fija de la intranet
    url = "http://reportes03/reports/report/IyD/Gestion/Funcionalidad"
    driver.get(url)

    WebDriverWait(driver, 30).until(
        EC.presence_of_element_located((By.TAG_NAME, "body"))
    )
    print("Página cargada correctamente.")
    time.sleep(3)

    # Cambiar al iframe del reporte
    iframe = driver.find_element(By.TAG_NAME, "iframe")
    driver.switch_to.frame(iframe)
    print("Cambiado al iframe del reporte.")

    # Buscar el campo de número de funcionalidad y asignar el valor
    numero_funcionalidad = "9107"
    input_funcionalidad = WebDriverWait(driver, 30).until(
        EC.visibility_of_element_located((By.ID, "ReportViewerControl_ctl04_ctl03_txtValue"))
    )
    driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", input_funcionalidad)
    driver.execute_script("arguments[0].focus();", input_funcionalidad)
    input_funcionalidad.click()
    input_funcionalidad.clear()
    for char in numero_funcionalidad:
        input_funcionalidad.send_keys(char)
        time.sleep(0.1)
    print(f"Número de funcionalidad {numero_funcionalidad} asignado.")
    time.sleep(0.5)

    # Click en el botón "Ver informe"
    ver_informe_btn = WebDriverWait(driver, 20).until(
        EC.element_to_be_clickable((By.ID, "ReportViewerControl_ctl04_ctl00"))
    )
    ver_informe_btn.click()
    print("Botón 'Ver informe' clickeado dentro del iframe.")
    time.sleep(15)

    # Buscar los divs de la columna "Numero" y filtrar requerimientos reales
    requerimiento_nros = []
    requerimiento_links = []
    divs = driver.find_elements(By.XPATH, "//div[contains(@style, 'width:20.22mm;min-width: 20.22mm;')]")
    for div in divs:
        try:
            # Si hay <a>, filtrar por href
            a_tag = div.find_element(By.TAG_NAME, "a")
            href = a_tag.get_attribute("href")
            nro = a_tag.text.strip()
            if href and "/Gestion/Requerimiento" in href and nro.isdigit():
                requerimiento_nros.append(nro)
                requerimiento_links.append(href)
        except Exception:
            # Si no hay <a>, tomar el texto solo si es número (fallback)
            nro = div.text.strip()
            if nro.isdigit():
                # Construir el link manualmente como antes
                href = f"http://reportes03/reports/report/IyD/Gestion/Requerimiento?TipoRequerimiento=1&Numero={nro}"
                requerimiento_nros.append(nro)
                requerimiento_links.append(href)
    print(f"Requerimientos encontrados: {requerimiento_nros}")

    # Guardar el HTML principal
    with open(html_path, "w", encoding="utf-8") as f:
        f.write(driver.page_source)
    print(f"HTML del informe guardado en {html_path}")

    # Crear y guardar el docx con la info de cada requerimiento
    document = Document()
    # === NUEVO: Extraer datos de la funcionalidad del HTML principal ===
    funcionalidad_campos = {
        'Funcionalidad nro.': numero_funcionalidad,
        'Nombre': '',
        'Descripción': '',
        'Producto': '',
        'Equipo': '',
        'Fecha alta': ''
    }
    # Analizar el HTML principal (ya descargado en html_path)
    with open(html_path, 'r', encoding='utf-8') as f:
        soup_func = BeautifulSoup(f, 'html.parser')
    for row in soup_func.find_all('tr'):
        celdas = row.find_all(['td', 'th'])
        if len(celdas) >= 2:
            label = celdas[0].get_text(strip=True).lower()
            valor = celdas[1].get_text(strip=True)
            if 'nombre' in label:
                funcionalidad_campos['Nombre'] = valor
            elif 'descrip' in label:
                funcionalidad_campos['Descripción'] = valor
            elif 'producto' in label:
                funcionalidad_campos['Producto'] = valor
            elif 'equipo' in label:
                funcionalidad_campos['Equipo'] = valor
            elif 'fecha alta' in label:
                funcionalidad_campos['Fecha alta'] = valor
    # Insertar los campos de funcionalidad al principio del documento
    document.add_paragraph(f"Funcionalidad nro.{funcionalidad_campos['Funcionalidad nro.']}")
    document.add_paragraph(f"• Nombre: {funcionalidad_campos['Nombre']}")
    document.add_paragraph(f"• Descripción: {funcionalidad_campos['Descripción']}")
    document.add_paragraph(f"• Producto: {funcionalidad_campos['Producto']}")
    document.add_paragraph(f"• Equipo: {funcionalidad_campos['Equipo']}")
    document.add_paragraph(f"• Fecha alta: {funcionalidad_campos['Fecha alta']}")
    document.add_paragraph("")

    for idx, nro in enumerate(requerimiento_nros):
        req_url = requerimiento_links[idx]
        driver.execute_script(f"window.open('{req_url}', '_blank');")
        driver.switch_to.window(driver.window_handles[-1])
        print(f"Requerimiento {nro} abierto en nueva pestaña.")
        time.sleep(5)
        html_req = driver.page_source
        soup = BeautifulSoup(html_req, 'html.parser')

        # Cambiar al iframe del requerimiento antes de extraer datos
        try:
            iframe_req = driver.find_element(By.TAG_NAME, "iframe")
            driver.switch_to.frame(iframe_req)
            print("Cambiado al iframe del requerimiento.")
            time.sleep(2)
            html_req = driver.page_source
            soup = BeautifulSoup(html_req, 'html.parser')
        except Exception as e:
            print(f"No se pudo cambiar al iframe del requerimiento: {e}")
            soup = BeautifulSoup(html_req, 'html.parser')
        finally:
            # Volver al contexto principal para cerrar la pestaña después
            driver.switch_to.default_content()

        # Nueva lógica para extraer campos por filas de tabla (dentro del iframe)
        campos = {
            'Fecha alta': '',
            'Título': '',
            'Necesidad': '',
            'Implementación sugerida': '',
            'Cliente': '',
            'Relevamientos asociados': '',
            'Documento número': '',
            'Link': '',
            'Interacciones relacionadas': ''
        }
        # Resetear variables para cada requerimiento
        documento_num = ''
        documento_link = ''
        incidente_num = ''
        incidente_link = ''
        for row in soup.find_all('tr'):
            celdas = row.find_all(['td', 'th'])
            if len(celdas) >= 2:
                label = celdas[0].get_text(strip=True)
                valor = celdas[1].get_text(strip=True)
                # Ajuste para los campos según el HTML real
                if 'fecha alta' in label.lower():
                    campos['Fecha alta'] = valor
                elif 'nombre' in label.lower():
                    campos['Título'] = valor
                elif 'necesidad' in label.lower():
                    campos['Necesidad'] = valor
                elif 'implem. sugerida' in label.lower():
                    campos['Implementación sugerida'] = valor
                elif 'cliente' in label.lower():
                    campos['Cliente'] += valor  # Puede haber dos filas: número y nombre
                elif 'requerido por' in label.lower():
                    campos['Relevamientos asociados'] = valor
                elif 'documento' == label.lower():
                    # Extraer solo el número de documento
                    doc_num_match = re.search(r'(\d+)', valor)
                    if doc_num_match:
                        documento_num = doc_num_match.group(1)
                        campos['Documento número'] = documento_num
                    else:
                        campos['Documento número'] = 'no hay'
                    # Buscar link del documento
                    link_tag = celdas[1].find('a', href=True)
                    if link_tag:
                        href = link_tag['href']
                        js_match = re.search(r"window.open\('([^']+)'", href)
                        if js_match:
                            documento_link = js_match.group(1)
                        else:
                            documento_link = href
                elif 'observación' in label.lower():
                    campos['Interacciones relacionadas'] = valor
        # Buscar incidente en la tabla (fuera de los tr principales, por divs con link)
        incidente_tag = soup.find('a', href=True, text=re.compile(r'\d{6,}'))
        if incidente_tag:
            incidente_num = incidente_tag.get_text(strip=True)
            href = incidente_tag['href']
            js_match = re.search(r"window.open\('([^']+)'", href)
            if js_match:
                incidente_link = js_match.group(1)
            else:
                incidente_link = href
        # Formatear el campo Cliente para agregar el separador si corresponde
        if campos['Cliente']:
            match = re.match(r"(\d+)(.*)", campos['Cliente'])
            if match:
                nro_cliente = match.group(1).strip()
                desc = match.group(2).strip()
                if desc:
                    campos['Cliente'] = f"{nro_cliente} - {desc}"
                else:
                    campos['Cliente'] = nro_cliente
        # Asignar 'no hay' a los campos vacíos excepto Link e Interacciones relacionadas
        for k in ['Fecha alta', 'Título', 'Necesidad', 'Implementación sugerida', 'Cliente', 'Relevamientos asociados', 'Documento número']:
            if not campos[k]:
                campos[k] = 'no hay'
        # Guardar los datos en el orden correcto
        document.add_paragraph(f"Requerimiento nro.{nro}")
        document.add_paragraph(f"• Fecha alta: {campos['Fecha alta']}")
        document.add_paragraph(f"• Título: {campos['Título']}")
        document.add_paragraph(f"• Necesidad: {campos['Necesidad']}")
        document.add_paragraph(f"• Implementación sugerida: {campos['Implementación sugerida']}")
        document.add_paragraph(f"• Cliente: {campos['Cliente']}")
        document.add_paragraph(f"• Relevamientos asociados: {campos['Relevamientos asociados']}")
        document.add_paragraph(f"• Documento número: {campos['Documento número']}")
        # Link: si hay documento_link y documento_num, poner hipervínculo, si no, 'no hay'
        if documento_link and documento_num:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            p = document.add_paragraph()
            run = p.add_run(f"• Link: ")
            r_id = document.part.relate_to(documento_link, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)
            new_run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            color = OxmlElement('w:color')
            color.set(qn('w:val'), '0000FF')
            rPr.append(color)
            u = OxmlElement('w:u')
            u.set(qn('w:val'), 'single')
            rPr.append(u)
            new_run.append(rPr)
            t = OxmlElement('w:t')
            t.text = documento_num
            new_run.append(t)
            hyperlink.append(new_run)
            p._p.append(hyperlink)
        else:
            document.add_paragraph(f"• Link: no hay")
        # Interacciones relacionadas siempre al final
        if incidente_num and incidente_link:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            p = document.add_paragraph()
            run = p.add_run(f"• Interacciones relacionadas: Incidente ")
            r_id = document.part.relate_to(incidente_link, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)
            new_run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            color = OxmlElement('w:color')
            color.set(qn('w:val'), '0000FF')
            rPr.append(color)
            u = OxmlElement('w:u')
            u.set(qn('w:val'), 'single')
            rPr.append(u)
            new_run.append(rPr)
            t = OxmlElement('w:t')
            t.text = incidente_num
            new_run.append(t)
            hyperlink.append(new_run)
            p._p.append(hyperlink)
        else:
            document.add_paragraph(f"• Interacciones relacionadas: no hay")
        document.add_paragraph("")
        driver.close()
        driver.switch_to.window(driver.window_handles[0])
    document.save(docx_path)
    print(f"Documento Word guardado en {docx_path}")

    print("Proceso finalizado. Cierra el navegador.")
finally:
    driver.quit()
