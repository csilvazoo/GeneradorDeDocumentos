# Versión 4: Elimina el requerimiento vacío de la plantilla y agrega datos de la funcionalidad con estilos y ubicación correctos.
import time
import re
import os
from selenium import webdriver
from selenium.webdriver.edge.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from docx import Document
from bs4 import BeautifulSoup
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

driver_path = None
options = Options()
options.add_argument("--start-maximized")
driver = webdriver.Edge(options=options)

try:
    base_path = os.path.dirname(os.path.abspath(__file__))
    html_path = os.path.join(base_path, "reporte_funcionalidad.html")
    docx_path = os.path.join(base_path, "requerimientos_v4.docx")
    plantilla_path = os.path.join(base_path, "plantillaAFU2505.docx")

    url = "http://reportes03/reports/report/IyD/Gestion/Funcionalidad"
    driver.get(url)
    WebDriverWait(driver, 30).until(EC.presence_of_element_located((By.TAG_NAME, "body")))
    time.sleep(3)
    iframe = driver.find_element(By.TAG_NAME, "iframe")
    driver.switch_to.frame(iframe)
    numero_funcionalidad = "9107"
    input_funcionalidad = WebDriverWait(driver, 30).until(
        EC.visibility_of_element_located((By.ID, "ReportViewerControl_ctl04_ctl03_txtValue"))
    )
    driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", input_funcionalidad)
    driver.execute_script("arguments[0].focus();", input_funcionalidad)
    input_funcionalidad.click()
    input_funcionalidad.clear()
    for char in numero_funcionalidad:
        input_funcionalidad.send_keys(char)
        time.sleep(0.1)
    time.sleep(0.5)
    ver_informe_btn = WebDriverWait(driver, 20).until(
        EC.element_to_be_clickable((By.ID, "ReportViewerControl_ctl04_ctl00"))
    )
    ver_informe_btn.click()
    time.sleep(15)

    divs = driver.find_elements(By.XPATH, "//div[contains(@style, 'width:20.22mm;min-width: 20.22mm;') or contains(@style, 'width:23.24mm;min-width: 23.24mm;')]")
    requerimiento_nros = []
    requerimiento_links = []
    after_req_cliente = False
    for div in divs:
        text = div.text.strip()
        if text == "Req. Cliente":
            after_req_cliente = True
            continue
        if after_req_cliente:
            if text.isdigit() and len(text) < 7:
                nro = text
                href = f"http://reportes03/reports/report/IyD/Gestion/Requerimiento?TipoRequerimiento=1&Numero={nro}"
                requerimiento_nros.append(nro)
                requerimiento_links.append(href)
            elif text.isdigit() and len(text) >= 7:
                continue
            elif text == "Req. Cliente":
                continue
    with open(html_path, "w", encoding="utf-8") as f:
        f.write(driver.page_source)

    document = Document(plantilla_path)
    # Eliminar el bloque de requerimiento vacío que deja la plantilla
    paras_to_remove = []
    found_req = False
    for i, p in enumerate(document.paragraphs):
        if p.text.strip().lower().startswith('requerimiento nro.'):
            # Marcar todos los párrafos hasta el siguiente encabezado o hasta 10 después
            for j in range(i, min(i+10, len(document.paragraphs))):
                if document.paragraphs[j].text.strip().lower().startswith('requerimiento nro.') and j != i:
                    break
                paras_to_remove.append(document.paragraphs[j])
            found_req = True
            break
    for p in paras_to_remove:
        p._element.getparent().remove(p._element)

    def insert_paragraph_after(paragraph, text, style=None):
        new_paragraph = paragraph._parent.add_paragraph(text, style=style)
        paragraph._element.addnext(new_paragraph._element)
        return new_paragraph

    # Buscar el párrafo con el título "Requerimientos" y guardar referencia
    req_paragraph = None
    for p in document.paragraphs:
        if p.text.strip().lower() == "requerimientos":
            req_paragraph = p
            break
    if req_paragraph is None:
        req_paragraph = document.paragraphs[-1]

    # Insertar los requerimientos después del título "Requerimientos"
    last_paragraph = req_paragraph
    for idx, nro in enumerate(requerimiento_nros):
        req_url = requerimiento_links[idx]
        driver.execute_script(f"window.open('{req_url}', '_blank');")
        driver.switch_to.window(driver.window_handles[-1])
        time.sleep(5)
        html_req = driver.page_source
        soup = BeautifulSoup(html_req, 'html.parser')
        try:
            iframe_req = driver.find_element(By.TAG_NAME, "iframe")
            driver.switch_to.frame(iframe_req)
            time.sleep(2)
            html_req = driver.page_source
            soup = BeautifulSoup(html_req, 'html.parser')
        except Exception:
            soup = BeautifulSoup(html_req, 'html.parser')
        finally:
            driver.switch_to.default_content()
        campos = {
            'Fecha alta': '',
            'Título': '',
            'Necesidad': '',
            'Implementación sugerida': '',
            'Cliente': '',
            'Relevamientos asociados': '',
            'Documento número': '',
            'Link': '',
            'Interacciones relacionadas': ''
        }
        documento_num = ''
        documento_link = ''
        incidente_num = ''
        incidente_link = ''
        for row in soup.find_all('tr'):
            celdas = row.find_all(['td', 'th'])
            if len(celdas) >= 2:
                label = celdas[0].get_text(strip=True)
                valor = celdas[1].get_text(strip=True)
                if 'fecha alta' in label.lower():
                    campos['Fecha alta'] = valor
                elif 'nombre' in label.lower():
                    campos['Título'] = valor
                elif 'necesidad' in label.lower():
                    campos['Necesidad'] = valor
                elif 'implem. sugerida' in label.lower():
                    campos['Implementación sugerida'] = valor
                elif 'cliente' in label.lower():
                    campos['Cliente'] += valor
                elif 'requerido por' in label.lower():
                    campos['Relevamientos asociados'] = valor
                elif 'documento' == label.lower():
                    doc_num_match = re.search(r'(\d+)', valor)
                    if doc_num_match:
                        documento_num = doc_num_match.group(1)
                        campos['Documento número'] = documento_num
                    else:
                        campos['Documento número'] = ''
                    link_tag = celdas[1].find('a', href=True)
                    if link_tag:
                        href = link_tag['href']
                        js_match = re.search(r"window.open\('([^']+)'", href)
                        if js_match:
                            documento_link = js_match.group(1)
                        else:
                            documento_link = href
                elif 'observación' in label.lower():
                    campos['Interacciones relacionadas'] = valor
        incidente_tag = soup.find('a', href=True, string=re.compile(r'\d{6,}'))
        if incidente_tag:
            incidente_num = incidente_tag.get_text(strip=True)
            href = incidente_tag['href']
            js_match = re.search(r"window.open\('([^']+)'", href)
            if js_match:
                incidente_link = js_match.group(1)
            else:
                incidente_link = href
        if campos['Cliente']:
            match = re.match(r"(\d+)(.*)", campos['Cliente'])
            if match:
                nro_cliente = match.group(1).strip()
                desc = match.group(2).strip()
                if desc:
                    campos['Cliente'] = f"{nro_cliente} - {desc}"
                else:
                    campos['Cliente'] = nro_cliente
        if documento_num:
            campos['Relevamientos asociados'] = ''
            campos['Documento número'] = documento_num
        else:
            campos['Relevamientos asociados'] = 'no hay'
            campos['Documento número'] = 'no hay'
            documento_link = ''
        for k in ['Fecha alta', 'Título', 'Necesidad', 'Implementación sugerida', 'Cliente']:
            if not campos[k]:
                campos[k] = 'no hay'
        p = insert_paragraph_after(last_paragraph, f"Requerimiento nro.{nro}", style="Encabezados del documento")
        p = insert_paragraph_after(p, f"Fecha alta: {campos['Fecha alta']}", style="List Bullet")
        p = insert_paragraph_after(p, f"Título: {campos['Título']}", style="List Bullet")
        p = insert_paragraph_after(p, f"Necesidad: {campos['Necesidad']}", style="List Bullet")
        p = insert_paragraph_after(p, f"Implementación sugerida: {campos['Implementación sugerida']}", style="List Bullet")
        p = insert_paragraph_after(p, f"Cliente: {campos['Cliente']}", style="List Bullet")
        p = insert_paragraph_after(p, f"Relevamientos asociados: {campos['Relevamientos asociados']}", style="List Bullet")
        p = insert_paragraph_after(p, f"Documento número: {campos['Documento número']}", style="List Bullet 2")
        if documento_link and documento_num:
            p = insert_paragraph_after(p, f"Link: ", style="List Bullet 2")
            r_id = document.part.relate_to(documento_link, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)
            new_run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            color = OxmlElement('w:color')
            color.set(qn('w:val'), '0000FF')
            rPr.append(color)
            u = OxmlElement('w:u')
            u.set(qn('w:val'), 'single')
            rPr.append(u)
            new_run.append(rPr)
            t = OxmlElement('w:t')
            t.text = documento_num
            new_run.append(t)
            hyperlink.append(new_run)
            p._p.append(hyperlink)
        else:
            p = insert_paragraph_after(p, f"Link: no hay", style="List Bullet 2")
        if incidente_num and incidente_link:
            p = insert_paragraph_after(p, f"Interacciones relacionadas: ", style="List Bullet")
            p = insert_paragraph_after(p, f"Incidente ", style="List Bullet 2")
            r_id = document.part.relate_to(incidente_link, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)
            new_run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            color = OxmlElement('w:color')
            color.set(qn('w:val'), '0000FF')
            rPr.append(color)
            u = OxmlElement('w:u')
            u.set(qn('w:val'), 'single')
            rPr.append(u)
            new_run.append(rPr)
            t = OxmlElement('w:t')
            t.text = incidente_num
            new_run.append(t)
            hyperlink.append(new_run)
            p._p.append(hyperlink)
        else:
            p = insert_paragraph_after(p, f"Interacciones relacionadas: no hay", style="List Bullet")
        last_paragraph = p

    # Insertar los datos de la funcionalidad después del último requerimiento y antes de "Notas aclaratorias"
    # Buscar el párrafo con "Notas aclaratorias"
    notas_paragraph = None
    for p in document.paragraphs:
        if p.text.strip().lower().startswith("notas aclaratorias"):
            notas_paragraph = p
            break
    if notas_paragraph is None:
        notas_paragraph = document.paragraphs[-1]

    # Extraer datos reales de la funcionalidad del primer HTML
    with open(html_path, 'r', encoding='utf-8') as f:
        html_func = f.read()
    soup_func = BeautifulSoup(html_func, 'html.parser')
    campos_func = {
        'Funcionalidad nro.': numero_funcionalidad,
        'Nombre:': '',
        'Descripción:': '',
        'Producto:': '',
        'Equipo:': '',
        'Fecha alta:': ''
    }
    for row in soup_func.find_all('tr'):
        celdas = row.find_all(['td', 'th'])
        if len(celdas) >= 2:
            label = celdas[0].get_text(strip=True).lower()
            valor = celdas[1].get_text(strip=True)
            if 'nombre' in label:
                campos_func['Nombre:'] = valor
            elif 'descripción' in label:
                campos_func['Descripción:'] = valor
            elif 'producto' in label:
                campos_func['Producto:'] = valor
            elif 'equipo' in label:
                campos_func['Equipo:'] = valor
            elif 'fecha alta' in label:
                campos_func['Fecha alta:'] = valor
    for k in campos_func:
        if not campos_func[k]:
            campos_func[k] = '(no hay)'
    # Insertar los datos de la funcionalidad justo antes de "Notas aclaratorias"
    insert_after = last_paragraph
    for i, p in enumerate(document.paragraphs):
        if p is last_paragraph:
            insert_after = p
            break
    for label, value in campos_func.items():
        insert_after = insert_paragraph_after(insert_after, f"{label} {value}", style="List Bullet")

    document.save(docx_path)
    print(f"Documento Word guardado en {docx_path}")
finally:
    driver.quit()
