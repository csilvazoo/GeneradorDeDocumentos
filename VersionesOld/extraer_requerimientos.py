import bs4
from docx import Document
from docx.shared import Pt

# Leer el HTML guardado
with open(r"c:\Users\csilva\Desktop\reporte_funcionalidad.html", encoding="utf-8") as f:
    html = f.read()

soup = bs4.BeautifulSoup(html, "html.parser")

# Crear el documento Word
doc = Document()

def add_req(doc, req):
    doc.add_paragraph(req, style='List Bullet')

# Buscar los requerimientos (ajustar según la estructura real)
# Buscamos por texto clave y agrupamos los bloques
text = soup.get_text("\n")
lines = [line.strip() for line in text.splitlines() if line.strip()]

reqs = []
req = []
for line in lines:
    if line.lower().startswith("req. cliente") or line.lower().startswith("requerimiento nro."):
        if req:
            reqs.append(req)
            req = []
    req.append(line)
if req:
    reqs.append(req)

for req_block in reqs:
    for line in req_block:
        doc.add_paragraph(line, style='List Bullet')
    doc.add_paragraph("")

# Guardar el docx
doc.save(r"c:\Users\csilva\Desktop\requerimientos_funcionalidad.docx")
print("Documento Word generado: requerimientos_funcionalidad.docx")
