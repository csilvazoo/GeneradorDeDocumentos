import time
from selenium import webdriver
from selenium.webdriver.edge.service import Service as EdgeService
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.edge.options import Options
from bs4 import BeautifulSoup
from docx import Document
import os

# Configuración de Edge
options = Options()
options.add_argument("--start-maximized")
driver = webdriver.Edge(options=options)

try:
    url = "http://reportes03/reports/report/IyD/Gestion/Funcionalidad"
    driver.get(url)
    WebDriverWait(driver, 30).until(
        EC.presence_of_element_located((By.TAG_NAME, "body"))
    )
    print("Página cargada correctamente.")
    time.sleep(3)

    # Cambiar al iframe del reporte
    iframe = driver.find_element(By.TAG_NAME, "iframe")
    driver.switch_to.frame(iframe)
    print("Cambiado al iframe del reporte.")

    # Buscar el campo de número de funcionalidad y asignar el valor
    numero_funcionalidad = "9107"
    input_funcionalidad = WebDriverWait(driver, 30).until(
        EC.visibility_of_element_located((By.ID, "ReportViewerControl_ctl04_ctl03_txtValue"))
    )
    driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", input_funcionalidad)
    driver.execute_script("arguments[0].focus();", input_funcionalidad)
    input_funcionalidad.click()
    input_funcionalidad.clear()
    for char in numero_funcionalidad:
        input_funcionalidad.send_keys(char)
        time.sleep(0.1)
    print(f"Número de funcionalidad {numero_funcionalidad} asignado.")
    time.sleep(0.5)

    # Click en el botón "Ver informe"
    ver_informe_btn = WebDriverWait(driver, 20).until(
        EC.element_to_be_clickable((By.ID, "ReportViewerControl_ctl04_ctl00"))
    )
    ver_informe_btn.click()
    print("Botón 'Ver informe' clickeado dentro del iframe.")
    time.sleep(10)

    # Extraer los datos de la funcionalidad de manera precisa
    html_func = driver.page_source
    soup_func = BeautifulSoup(html_func, 'html.parser')
    campos_func = {
        'Nombre': '',
        'Descripción': '',
        'Producto': '',
        'Equipo': '',
        'Fecha alta': ''
    }
    # Buscar solo la celda exacta y tomar la siguiente o la siguiente fila/tabla si es necesario
    for row_idx, row in enumerate(soup_func.find_all('tr')):
        celdas = row.find_all(['td', 'th'])
        for idx, celda in enumerate(celdas):
            texto = celda.get_text(strip=True).lower().replace(':','').replace('á','a').replace('é','e').replace('í','i').replace('ó','o').replace('ú','u')
            # Nombre y Descripción funcionan bien con la celda siguiente
            if texto == 'nombre' and not campos_func['Nombre'] and idx+1 < len(celdas):
                campos_func['Nombre'] = celdas[idx+1].get_text(strip=True)
            elif texto == 'descripcion' and not campos_func['Descripción'] and idx+1 < len(celdas):
                campos_func['Descripción'] = celdas[idx+1].get_text(strip=True)
            # Producto: buscar en la celda siguiente, y si está vacía, buscar en la siguiente fila, misma columna
            elif texto in ['producto', 'productos'] and not campos_func['Producto']:
                valor = ''
                # 1. Celda siguiente
                if idx+1 < len(celdas):
                    valor = celdas[idx+1].get_text(strip=True)
                    # Si la celda tiene una tabla, buscar el primer texto no vacío
                    if not valor:
                        inner_table = celdas[idx+1].find('table')
                        if inner_table:
                            valor = inner_table.get_text(strip=True)
                # 2. Si sigue vacío, buscar en la siguiente fila, misma columna
                if not valor:
                    filas = soup_func.find_all('tr')
                    if row_idx+1 < len(filas):
                        next_row = filas[row_idx+1]
                        next_celdas = next_row.find_all(['td', 'th'])
                        if idx < len(next_celdas):
                            valor = next_celdas[idx].get_text(strip=True)
                            if not valor:
                                inner_table = next_celdas[idx].find('table')
                                if inner_table:
                                    valor = inner_table.get_text(strip=True)
                campos_func['Producto'] = valor
            # Equipo: igual que antes
            elif texto == 'equipo' and not campos_func['Equipo'] and idx+1 < len(celdas):
                campos_func['Equipo'] = celdas[idx+1].get_text(strip=True)
            # Fecha alta: buscar en la celda siguiente, y si está vacía, buscar en la siguiente fila, misma columna
            elif texto == 'fecha alta' and not campos_func['Fecha alta']:
                valor = ''
                if idx+1 < len(celdas):
                    valor = celdas[idx+1].get_text(strip=True)
                    if not valor:
                        inner_table = celdas[idx+1].find('table')
                        if inner_table:
                            valor = inner_table.get_text(strip=True)
                if not valor:
                    filas = soup_func.find_all('tr')
                    if row_idx+1 < len(filas):
                        next_row = filas[row_idx+1]
                        next_celdas = next_row.find_all(['td', 'th'])
                        if idx < len(next_celdas):
                            valor = next_celdas[idx].get_text(strip=True)
                            if not valor:
                                inner_table = next_celdas[idx].find('table')
                                if inner_table:
                                    valor = inner_table.get_text(strip=True)
                campos_func['Fecha alta'] = valor
    # Fallback: si algún campo sigue vacío, dejar 'no hay'
    for k in campos_func:
        if not campos_func[k]:
            campos_func[k] = 'no hay'

    # Guardar los datos en un Word
    docx_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "word.docx")
    document = Document()
    document.add_heading(f"Funcionalidad nro. {numero_funcionalidad}", level=2)
    document.add_paragraph(f"Nombre: {campos_func['Nombre']}")
    document.add_paragraph(f"Descripción: {campos_func['Descripción']}")
    document.add_paragraph(f"Producto: {campos_func['Producto']}")
    document.add_paragraph(f"Equipo: {campos_func['Equipo']}")
    document.add_paragraph(f"Fecha alta: {campos_func['Fecha alta']}")
    document.save(docx_path)
    print(f"Datos guardados en {docx_path}")
finally:
    driver.quit()
