import sys
import threading
import tkinter as tk
from tkinter import ttk, messagebox
from tkinter.scrolledtext import ScrolledText
import queue
import time
import os
import getpass
from tkinter import filedialog

log_queue = queue.Queue()

def run_script(numero_funcionalidad, ver_explorador, log_queue, docx_path):
    import re
    import os
    import time
    from selenium import webdriver
    from selenium.webdriver.common.by import By
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    from selenium.webdriver.edge.options import Options
    from selenium.webdriver.edge.service import Service as EdgeService
    from docx import Document
    from bs4 import BeautifulSoup

    def log(msg):
        log_queue.put(msg)

    options = Options()
    options.add_argument("--start-maximized")
    if not ver_explorador:
        options.add_argument("--headless")
    # Buscar msedgedriver.exe en el mismo directorio (soporta PyInstaller)
    if getattr(sys, 'frozen', False):
        driver_path = os.path.join(sys._MEIPASS, "msedgedriver.exe")
    else:
        driver_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "msedgedriver.exe")
    service = EdgeService(driver_path)
    driver = webdriver.Edge(service=service, options=options)
    try:
        base_path = os.path.dirname(os.path.abspath(__file__))
        plantilla_path = os.path.join(base_path, "plantillaAFU2505.docx")
        url = "http://reportes03/reports/report/IyD/Gestion/Funcionalidad"
        driver.get(url)
        WebDriverWait(driver, 30).until(
            EC.presence_of_element_located((By.TAG_NAME, "body"))
        )
        log("Página cargada correctamente.")
        time.sleep(3)
        iframe = driver.find_element(By.TAG_NAME, "iframe")
        driver.switch_to.frame(iframe)
        log("Cambiado al iframe del reporte.")
        input_funcionalidad = WebDriverWait(driver, 30).until(
            EC.visibility_of_element_located((By.ID, "ReportViewerControl_ctl04_ctl03_txtValue"))
        )
        driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", input_funcionalidad)
        driver.execute_script("arguments[0].focus();", input_funcionalidad)
        input_funcionalidad.click()
        input_funcionalidad.clear()
        for char in numero_funcionalidad:
            input_funcionalidad.send_keys(char)
            time.sleep(0.1)
        log(f"Número de funcionalidad {numero_funcionalidad} asignado.")
        time.sleep(0.5)
        ver_informe_btn = WebDriverWait(driver, 20).until(
            EC.element_to_be_clickable((By.ID, "ReportViewerControl_ctl04_ctl00"))
        )
        ver_informe_btn.click()
        log("Botón 'Ver informe' clickeado dentro del iframe.")
        time.sleep(15)
        divs = driver.find_elements(By.XPATH, "//div[contains(@style, 'width:20.22mm;min-width: 20.22mm;') or contains(@style, 'width:23.24mm;min-width: 23.24mm;')]")
        requerimiento_nros = []
        requerimiento_links = []
        after_req_cliente = False
        for div in divs:
            text = div.text.strip()
            if text == "Req. Cliente":
                after_req_cliente = True
                continue
            if after_req_cliente:
                if text.isdigit() and len(text) < 7:
                    nro = text
                    href = f"http://reportes03/reports/report/IyD/Gestion/Requerimiento?TipoRequerimiento=1&Numero={nro}"
                    requerimiento_nros.append(nro)
                    requerimiento_links.append(href)
                elif text.isdigit() and len(text) >= 7:
                    continue
                elif text == "Req. Cliente":
                    continue
        log(f"Requerimientos encontrados: {requerimiento_nros}")
        if not requerimiento_nros:
            log("No existe funcionalidad o no tiene requerimientos asociados.")
            try:
                import tkinter
                from tkinter import messagebox
                root = tkinter.Tk()
                root.withdraw()
                messagebox.showerror("Error", "No existe funcionalidad o no tiene requerimientos asociados.")
                root.destroy()
            except Exception:
                pass
            log("Proceso finalizado.")
            return
        document = Document(plantilla_path)
        def insert_paragraph_after(paragraph, text, style=None):
            new_paragraph = paragraph._parent.add_paragraph(text)
            paragraph._element.addnext(new_paragraph._element)
            if style:
                try:
                    new_paragraph.style = style
                except KeyError:
                    pass
            return new_paragraph
        req_paragraph = None
        for p in document.paragraphs:
            if p.text.strip().lower() == "requerimientos":
                req_paragraph = p
                break
        if req_paragraph is None:
            req_paragraph = document.paragraphs[-1]
        last_paragraph = req_paragraph
        for idx, nro in enumerate(requerimiento_nros):
            req_url = requerimiento_links[idx]
            driver.execute_script(f"window.open('{req_url}', '_blank');")
            driver.switch_to.window(driver.window_handles[-1])
            log(f"Requerimiento {nro} abierto en nueva pestaña.")
            time.sleep(7)
            html_req = driver.page_source
            soup = BeautifulSoup(html_req, 'html.parser')
            try:
                iframe_req = driver.find_element(By.TAG_NAME, "iframe")
                driver.switch_to.frame(iframe_req)
                log("Cambiado al iframe del requerimiento.")
                time.sleep(2)
                html_req = driver.page_source
                soup = BeautifulSoup(html_req, 'html.parser')
            except Exception as e:
                log(f"No se pudo cambiar al iframe del requerimiento: {e}")
                soup = BeautifulSoup(html_req, 'html.parser')
            finally:              
                driver.switch_to.default_content()
            campos = {
                'Fecha alta': '',
                'Título': '',
                'Necesidad': '',
                'Implementación sugerida': '',
                'Cliente': '',
                'Relevamientos asociados': '',
                'Documento número': '',
                'Link': '',
                'Interacciones relacionadas': ''
            }
            documento_num = ''
            documento_link = ''
            incidente_num = ''
            incidente_link = ''
            for row in soup.find_all('tr'):
                celdas = row.find_all(['td', 'th'])
                if len(celdas) >= 2:
                    label = celdas[0].get_text(strip=True)
                    valor = celdas[1].get_text(strip=True)
                    if 'fecha alta' in label.lower():
                        campos['Fecha alta'] = valor
                    elif 'nombre' in label.lower():
                        campos['Título'] = valor
                    elif 'necesidad' in label.lower():
                        campos['Necesidad'] = valor
                    elif 'implem. sugerida' in label.lower():
                        campos['Implementación sugerida'] = valor
                    elif 'cliente' in label.lower():
                        campos['Cliente'] += valor
                    elif 'requerido por' in label.lower():
                        campos['Relevamientos asociados'] = valor
                    elif 'documento' == label.lower():
                        doc_num_match = re.search(r'(\d+)', valor)
                        if doc_num_match:
                            documento_num = doc_num_match.group(1)
                            campos['Documento número'] = documento_num
                        else:
                            campos['Documento número'] = ''
                        link_tag = celdas[1].find('a', href=True)
                        if link_tag:
                            href = link_tag['href']
                            js_match = re.search(r"window.open\('([^']+)'", href)
                            if js_match:
                                documento_link = js_match.group(1)
                            else:
                                documento_link = href
                    elif 'observación' in label.lower():
                        campos['Interacciones relacionadas'] = valor
            incidente_tag = soup.find('a', href=True, text=re.compile(r'\d{6,}'))
            if incidente_tag:
                incidente_num = incidente_tag.get_text(strip=True)
                href = incidente_tag['href']
                js_match = re.search(r"window.open\('([^']+)'", href)
                if js_match:
                    incidente_link = js_match.group(1)
                else:
                    incidente_link = href
            if campos['Cliente']:
                match = re.match(r"(\d+)(.*)", campos['Cliente'])
                if match:
                    nro_cliente = match.group(1).strip()
                    desc = match.group(2).strip()
                    if desc:
                        campos['Cliente'] = f"{nro_cliente} - {desc}"
                    else:
                        campos['Cliente'] = nro_cliente
            if documento_num:
                campos['Relevamientos asociados'] = ''
                campos['Documento número'] = documento_num
            else:
                campos['Relevamientos asociados'] = 'no hay'
                campos['Documento número'] = 'no hay'
                documento_link = ''
            for k in ['Fecha alta', 'Título', 'Necesidad', 'Implementación sugerida', 'Cliente']:
                if not campos[k]:
                    campos[k] = 'no hay'
            p = insert_paragraph_after(last_paragraph, f"Requerimiento nro.{nro}", style="toa heading")
            p = insert_paragraph_after(p, f"Fecha alta: {campos['Fecha alta']}", style="List Bullet")
            p = insert_paragraph_after(p, f"Título: {campos['Título']}", style="List Bullet")
            p = insert_paragraph_after(p, f"Necesidad: {campos['Necesidad']}", style="List Bullet")
            p = insert_paragraph_after(p, f"Implementación sugerida: {campos['Implementación sugerida']}", style="List Bullet")
            p = insert_paragraph_after(p, f"Cliente: {campos['Cliente']}", style="List Bullet")
            p = insert_paragraph_after(p, f"Relevamientos asociados: {campos['Relevamientos asociados']}", style="List Bullet")
            p = insert_paragraph_after(p, f"Documento número: {campos['Documento número']}", style="List Bullet 2")
            if documento_link and documento_num:
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn
                p = insert_paragraph_after(p, f"Link: ", style="List Bullet 2")
                r_id = document.part.relate_to(documento_link, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
                hyperlink = OxmlElement('w:hyperlink')
                hyperlink.set(qn('r:id'), r_id)
                new_run = OxmlElement('w:r')
                rPr = OxmlElement('w:rPr')
                rStyle = OxmlElement('w:rStyle')
                rStyle.set(qn('w:val'), 'Hyperlink')
                rPr.append(rStyle)
                color = OxmlElement('w:color')
                color.set(qn('w:val'), '#009430')
                rPr.append(color)
                u = OxmlElement('w:u')
                u.set(qn('w:val'), 'single')
                rPr.append(u)
                new_run.append(rPr)
                t = OxmlElement('w:t')
                t.text = documento_num
                new_run.append(t)
                hyperlink.append(new_run)
                p._p.append(hyperlink)
            else:
                p = insert_paragraph_after(p, f"Link: no hay", style="List Bullet 2")
            if incidente_num and incidente_link:
                p = insert_paragraph_after(p, f"Interacciones relacionadas: ", style="List Bullet")
                p = insert_paragraph_after(p, f"Incidente: ", style="List Bullet 2")
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn
                r_id = document.part.relate_to(incidente_link, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
                hyperlink = OxmlElement('w:hyperlink')
                hyperlink.set(qn('r:id'), r_id)
                new_run = OxmlElement('w:r')
                rPr = OxmlElement('w:rPr')
                rStyle = OxmlElement('w:rStyle')
                rStyle.set(qn('w:val'), 'Hyperlink')
                rPr.append(rStyle)
                color = OxmlElement('w:color')
                color.set(qn('w:val'), '#009430')
                rPr.append(color)
                u = OxmlElement('w:u')
                u.set(qn('w:val'), 'single')
                rPr.append(u)
                new_run.append(rPr)
                t = OxmlElement('w:t')
                t.text = incidente_num
                new_run.append(t)
                hyperlink.append(new_run)
                p._p.append(hyperlink)
            else:
                p = insert_paragraph_after(p, f"Interacciones relacionadas: no hay", style="List Bullet")
            last_paragraph = p
        # Eliminar el último bloque de requerimiento si está vacío (de atrás hacia adelante)
        def is_empty_req_block(paragraphs, start_idx, labels):
            end_idx = start_idx + 1
            while end_idx < len(paragraphs):
                text = paragraphs[end_idx].text.strip().lower()
                if text.startswith('requerimiento nro.') or text.startswith('funcionalidad nro.'):
                    break
                end_idx += 1
            for i in range(start_idx, end_idx):
                t = paragraphs[i].text.strip().lower()
                if t and not any(t.startswith(label) for label in labels):
                    return False, end_idx
                if ':' in t:
                    val = t.split(':',1)[1].strip()
                    if val and val != 'no hay':
                        return False, end_idx
            return True, end_idx
        labels = [
            'requerimiento nro.',
            'fecha alta',
            'título',
            'necesidad',
            'implementación sugerida',
            'cliente',
            'relevamientos asociados',
            'documento número',
            'link',
            'interacciones relacionadas',
            'incidente'
        ]
        for i in range(len(document.paragraphs)-1, -1, -1):
            t = document.paragraphs[i].text.strip().lower()
            if t.startswith('requerimiento nro.'):
                empty, end_idx = is_empty_req_block(document.paragraphs, i, labels)
                if empty:
                    for j in range(end_idx-1, i-1, -1):
                        document.paragraphs[j]._element.getparent().remove(document.paragraphs[j]._element)
                    break
        # --- AGREGAR DATOS DE LA FUNCIONALIDAD ---
        driver.switch_to.window(driver.window_handles[0])
        driver.switch_to.default_content()
        iframe = driver.find_element(By.TAG_NAME, "iframe")
        driver.switch_to.frame(iframe)
        html_func = driver.page_source
        soup_func = BeautifulSoup(html_func, 'html.parser')
        campos_func = {
            'Funcionalidad nro.': numero_funcionalidad,
            'Nombre': '',
            'Descripción': '',
            'Producto': '',
            'Equipo': '',
            'Fecha alta': ''
        }
        for row_idx, row in enumerate(soup_func.find_all('tr')):
            celdas = row.find_all(['td', 'th'])
            for idx, celda in enumerate(celdas):
                texto = celda.get_text(strip=True).lower().replace(':','').replace('á','a').replace('é','e').replace('í','i').replace('ó','o').replace('ú','u')
                if texto == 'nombre' and not campos_func['Nombre'] and idx+1 < len(celdas):
                    campos_func['Nombre'] = celdas[idx+1].get_text(strip=True)
                elif texto == 'descripcion' and not campos_func['Descripción'] and idx+1 < len(celdas):
                    campos_func['Descripción'] = celdas[idx+1].get_text(strip=True)
                elif texto in ['producto', 'productos'] and not campos_func['Producto']:
                    valor = ''
                    if idx+1 < len(celdas):
                        valor = celdas[idx+1].get_text(strip=True)
                        if not valor:
                            inner_table = celdas[idx+1].find('table')
                            if inner_table:
                                valor = inner_table.get_text(strip=True)
                    if not valor:
                        filas = soup_func.find_all('tr')
                        if row_idx+1 < len(filas):
                            next_row = filas[row_idx+1]
                            next_celdas = next_row.find_all(['td', 'th'])
                            if idx < len(next_celdas):
                                valor = next_celdas[idx].get_text(strip=True)
                                if not valor:
                                    inner_table = next_celdas[idx].find('table')
                                    if inner_table:
                                        valor = inner_table.get_text(strip=True)
                    campos_func['Producto'] = valor
                elif texto == 'equipo' and not campos_func['Equipo'] and idx+1 < len(celdas):
                    campos_func['Equipo'] = celdas[idx+1].get_text(strip=True)
                elif texto == 'fecha alta' and not campos_func['Fecha alta']:
                    valor = ''
                    if idx+1 < len(celdas):
                        valor = celdas[idx+1].get_text(strip=True)
                        if not valor:
                            inner_table = celdas[idx+1].find('table')
                            if inner_table:
                                valor = inner_table.get_text(strip=True)
                    if not valor:
                        filas = soup_func.find_all('tr')
                        if row_idx+1 < len(filas):
                            next_row = filas[row_idx+1]
                            next_celdas = next_row.find_all(['td', 'th'])
                            if idx < len(next_celdas):
                                valor = next_celdas[idx].get_text(strip=True)
                                if not valor:
                                    inner_table = next_celdas[idx].find('table')
                                    if inner_table:
                                        valor = inner_table.get_text(strip=True)
                    campos_func['Fecha alta'] = valor
        for k in campos_func:
            if not campos_func[k]:
                campos_func[k] = 'no hay'
        insert_idx = None
        for i, p in enumerate(document.paragraphs):
            if p.text.strip().lower().startswith('funcionalidad nro.'):
                insert_idx = i
                break
        if insert_idx is not None:
            end_idx = insert_idx + 1
            while end_idx < len(document.paragraphs):
                t = document.paragraphs[end_idx].text.strip().lower()
                if t.startswith('•') or any(x in t for x in ['nombre', 'descrip', 'producto', 'equipo', 'fecha alta']):
                    end_idx += 1
                else:
                    break
            for j in range(end_idx-1, insert_idx-1, -1):
                document.paragraphs[j]._element.getparent().remove(document.paragraphs[j]._element)
            p = document.paragraphs[insert_idx-1] if insert_idx > 0 else document.paragraphs[0]
        else:
            p = last_paragraph
        p = insert_paragraph_after(p, f"Funcionalidad nro.{campos_func['Funcionalidad nro.']}", style="Heading 2")
        p = insert_paragraph_after(p, f"Nombre: {campos_func['Nombre']}", style="List Bullet")
        p = insert_paragraph_after(p, f"Descripción: {campos_func['Descripción']}", style="List Bullet")
        p = insert_paragraph_after(p, f"Producto: {campos_func['Producto']}", style="List Bullet")
        p = insert_paragraph_after(p, f"Equipo: {campos_func['Equipo']}", style="List Bullet")
        p = insert_paragraph_after(p, f"Fecha alta: {campos_func['Fecha alta']}", style="List Bullet")
        last_paragraph = p
        document.save(docx_path)
        log(f"Documento Word guardado en {docx_path}")
        log("Proceso finalizado.")
    except Exception as e:
        log(f"ERROR: {e}")
        log("Proceso finalizado.")
    finally:
        driver.quit()

def main():
    root = tk.Tk()
    root.title("Generador de Documentos Funcionales")
    root.geometry("600x220")
    frm = ttk.Frame(root, padding=20)
    frm.pack(fill=tk.BOTH, expand=True)
    ttk.Label(frm, text="Funcionalidad:").grid(row=0, column=0, sticky=tk.W)
    funcionalidad_var = tk.StringVar()
    txt_funcionalidad = ttk.Entry(frm, textvariable=funcionalidad_var, width=20)
    txt_funcionalidad.grid(row=0, column=1, sticky=tk.W)
    ver_explorador_var = tk.BooleanVar(value=True)
    chk = ttk.Checkbutton(frm, text="Ver explorador", variable=ver_explorador_var)
    chk.grid(row=1, column=0, columnspan=2, sticky=tk.W)
    # Guardar como debajo de ver explorador
    user_documents = os.path.join(os.path.expanduser('~'), 'Documents')
    ruta_var = tk.StringVar(value=os.path.join(user_documents, "Funcionalidad.docx"))
    ttk.Label(frm, text="Guardar como:").grid(row=2, column=0, sticky=tk.W)
    entry_ruta = ttk.Entry(frm, textvariable=ruta_var, width=65)
    entry_ruta.grid(row=2, column=1, sticky=tk.W, columnspan=1)
    def seleccionar_ruta():
        ruta = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Documentos Word", "*.docx")],
            initialdir=user_documents,
            initialfile="Funcionalidad.docx",
            title="Guardar documento como"
        )
        if ruta:
            ruta_var.set(ruta)
    btn_examinar = ttk.Button(frm, text="Examinar...", command=seleccionar_ruta)
    btn_examinar.grid(row=2, column=1, sticky=tk.E)
    btn = ttk.Button(frm, text="Generar Documento")
    btn.grid(row=3, column=0, pady=10, sticky=tk.W)
    open_btn = ttk.Button(frm, text="Abrir Documento", state="disabled")
    open_btn.grid(row=3, column=1, pady=10, sticky=tk.E)
    
    def abrir_doc():
        ruta_docx = ruta_var.get()
        if os.path.exists(ruta_docx):
            try:
                os.startfile(ruta_docx)
            except Exception as e:
                messagebox.showerror("Error", f"No se pudo abrir el documento:\n{e}")
        else:
            messagebox.showerror("Error", "El archivo no existe en la ruta seleccionada.")
    open_btn.config(command=abrir_doc)

    progress = ttk.Progressbar(frm, mode="indeterminate")
    progress.grid(row=4, column=0, columnspan=2, sticky="ew", pady=5)
    status_label = ttk.Label(frm, text="", foreground="blue")
    status_label.grid(row=5, column=0, columnspan=2, pady=(0,5))
    # Label de éxito/error y botón abrir documento (más compactos)
    success_label = ttk.Label(frm, text="", foreground="green")
    success_label.grid(row=8, column=0, columnspan=2, pady=(5,0))
    ttk.Label(frm, text="Detalle de ejecución:").grid(row=5, column=0, columnspan=2, sticky=tk.W)
    # --- Detalle de ejecución (área de log) ---
    from tkinter.scrolledtext import ScrolledText

    # Label para el título del log
    log_label = ttk.Label(frm, text="Detalle de ejecución:")
    log_label.grid(row=5, column=0, columnspan=2, sticky=tk.W)

    # Área de log (ScrolledText)
    txt_log = ScrolledText(frm, height=15, width=100, state="disabled", bg="#ffffff", bd=0, highlightthickness=0)
    txt_log.grid(row=7, column=0, columnspan=2, sticky="nsew", pady=(0,10))
    txt_log.grid_remove()  # Oculto por defecto

    # Permitir expansión del área de log al redimensionar la ventana
    frm.rowconfigure(7, weight=1)
    frm.columnconfigure(1, weight=1)

    # Botón para mostrar/ocultar el log
    def toggle_log():
        if txt_log.winfo_viewable():
            txt_log.grid_remove()
            ver_detalle_btn.config(text="Ver detalle")
            root.geometry("600x220")
        else:
            txt_log.grid(row=7, column=0, columnspan=2, sticky="nsew", pady=(0,10))
            ver_detalle_btn.config(text="Ocultar detalle")
            root.geometry("600x400")
    ver_detalle_btn = ttk.Button(frm, text="Ver detalle", command=toggle_log)
    ver_detalle_btn.grid(row=6, column=0, columnspan=2, pady=(0,10))

    # Función para actualizar el log en tiempo real
    def update_log():
        while not log_queue.empty():
            msg = log_queue.get_nowait()
            txt_log.configure(state="normal")
            txt_log.insert(tk.END, msg + "\n")
            txt_log.see(tk.END)
            txt_log.configure(state="disabled")
        root.after(200, update_log)
    update_log()

    def on_run():
        num = funcionalidad_var.get().strip()
        if not num.isdigit():
            messagebox.showerror("Error", "Ingrese un número de funcionalidad válido.")
            return
        btn.config(state="disabled")
        progress.start()
        txt_log.configure(state="normal")
        txt_log.delete(1.0, tk.END)
        txt_log.configure(state="disabled")
        success_label.config(text="", foreground="green")
        open_btn.config(state="disabled")
        status_label.config(text="Generando Documento", foreground="blue")
        root.geometry("600x220")
        ruta_docx = ruta_var.get()
        def limpiar_label():
            success_label.config(text="")
        def task():
            run_script(num, ver_explorador_var.get(), log_queue, ruta_docx)
            progress.stop()
            btn.config(state="normal")
            log_content = ""
            try:
                with log_queue.mutex:
                    log_content = "\n".join(list(log_queue.queue))
            except Exception:
                pass
            if not log_content:
                log_content = txt_log.get("1.0", tk.END)
            # Al finalizar, ocultar "Generando Documento" y mostrar resultado en status_label
            status_label.config(text="")
            if os.path.exists(ruta_docx) and "No existe funcionalidad o no tiene requerimientos asociados." not in log_content:
                status_label.config(text="Documento generado exitosamente.", foreground="green")
                open_btn.config(state="normal")
            else:
                status_label.config(text="Documento NO generado.", foreground="red")
                open_btn.config(state="disabled")
                root.after(5000, lambda: status_label.config(text=""))
        threading.Thread(target=task, daemon=True).start()
    btn.config(command=on_run)
    # Ventana compacta al inicio
    root.geometry("600x220")
    root.mainloop()

if __name__ == "__main__":
    main()
